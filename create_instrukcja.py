#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Skrypt do generowania instrukcji obsługi systemu wag w formacie Word
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_instrukcja():
    doc = Document()

    # Ustawienia strony
    sections = doc.sections
    for section in sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # ========== STRONA TYTUŁOWA ==========
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading('INSTRUKCJA OBSŁUGI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('SYSTEM WAGI / PANEL PRODUKCJI', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Wersja dokumentu: 1.0').bold = True
    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run('Data: ________________')

    doc.add_page_break()

    # ========== SPIS TREŚCI ==========
    doc.add_heading('SPIS TREŚCI', level=1)

    toc_items = [
        '1. Uruchomienie systemu',
        '2. Obsługa wagi - pobieranie danych',
        '3. Automatyczne wyświetlanie wagi',
        '4. Wprowadzanie danych ręcznych',
        '5. Zapisywanie i zatwierdzanie',
        '6. Rozwiązywanie problemów',
    ]

    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ========== SEKCJA 1: URUCHOMIENIE ==========
    doc.add_heading('1. URUCHOMIENIE SYSTEMU', level=1)

    doc.add_paragraph(
        'Po włączeniu terminala/tabletu system automatycznie uruchomi aplikację. '
        'W przypadku wyłączonej aplikacji - kliknij ikonę programu na pulpicie.'
    )

    # Miejsce na screenshot
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[MIEJSCE NA SCREENSHOT - EKRAN GŁÓWNY]')
    run.bold = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Ramka na screenshot
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(14)
    cell_para = cell.paragraphs[0]
    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_para.add_run('\n\n\n\n\n[Wklej tutaj screenshot ekranu głównego]\n\n\n\n\n')
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.size = Pt(12)

    # Formatowanie tabeli jako ramki
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(50)
            cell.paragraphs[0].paragraph_format.space_after = Pt(50)

    doc.add_paragraph()

    doc.add_page_break()

    # ========== SEKCJA 2: POBIERANIE DANYCH Z WAGI ==========
    doc.add_heading('2. OBSŁUGA WAGI - POBIERANIE DANYCH', level=1)

    doc.add_heading('2.1 Automatyczne pobieranie przez dotyk', level=2)

    important = doc.add_paragraph()
    run = important.add_run('WAŻNE: ')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    important.add_run(
        'Aby pobrać wagę z wyświetlacza wystarczy kliknąć palcem w pole wagi na ekranie. '
        'System automatycznie pobierze aktualną wartość z wagi.'
    )

    doc.add_paragraph()

    steps = doc.add_paragraph()
    steps.add_run('Kroki do wykonania:').bold = True

    step_list = [
        'Upewnij się, że towar jest na wadze i waga się ustabilizowała',
        'Kliknij palcem w pole "Waga" na ekranie dotykowym',
        'Wartość zostanie automatycznie pobrana i wyświetlona',
        'Zweryfikuj czy wartość jest prawidłowa',
    ]

    for i, step in enumerate(step_list, 1):
        p = doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_paragraph()

    # Miejsce na screenshot
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[MIEJSCE NA SCREENSHOT - POLE WAGI]')
    run.bold = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell_para = cell.paragraphs[0]
    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_para.add_run('\n\n\n\n\n[Wklej tutaj screenshot pola wagi]\n\n\n\n\n')
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.size = Pt(12)

    doc.add_paragraph()

    doc.add_page_break()

    # ========== SEKCJA 3: AUTOMATYCZNE WYŚWIETLANIE ==========
    doc.add_heading('3. AUTOMATYCZNE WYŚWIETLANIE WAGI', level=1)

    doc.add_paragraph(
        'System może być skonfigurowany do automatycznego wyświetlania wartości wagi. '
        'W tym trybie wartość z wagi aktualizuje się na bieżąco bez konieczności klikania.'
    )

    doc.add_paragraph()

    auto_info = doc.add_paragraph()
    run = auto_info.add_run('Tryb automatyczny:')
    run.bold = True

    auto_features = [
        'Waga wyświetla się automatycznie po położeniu towaru',
        'Aktualizacja następuje co kilka sekund',
        'Nie trzeba klikać w pole - wartość sama się pojawia',
        'Gdy waga się ustabilizuje, wartość zostanie zatwierdzona',
    ]

    for feature in auto_features:
        doc.add_paragraph(f'• {feature}')

    doc.add_paragraph()

    # Miejsce na screenshot
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[MIEJSCE NA SCREENSHOT - TRYB AUTOMATYCZNY]')
    run.bold = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell_para = cell.paragraphs[0]
    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_para.add_run('\n\n\n\n\n[Wklej tutaj screenshot trybu automatycznego]\n\n\n\n\n')
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.size = Pt(12)

    doc.add_paragraph()

    doc.add_page_break()

    # ========== SEKCJA 4: DANE RĘCZNE ==========
    doc.add_heading('4. WPROWADZANIE DANYCH RĘCZNYCH', level=1)

    doc.add_paragraph(
        'W przypadku awarii automatycznego pobierania lub potrzeby korekty, '
        'można wprowadzić wagę ręcznie.'
    )

    doc.add_paragraph()

    manual_steps = doc.add_paragraph()
    manual_steps.add_run('Wprowadzanie ręczne:').bold = True

    manual_list = [
        'Kliknij dwukrotnie w pole wagi',
        'Pojawi się klawiatura numeryczna',
        'Wprowadź wartość wagi',
        'Zatwierdź przyciskiem OK lub Enter',
    ]

    for i, step in enumerate(manual_list, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_paragraph()

    # Miejsce na screenshot
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[MIEJSCE NA SCREENSHOT - KLAWIATURA NUMERYCZNA]')
    run.bold = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell_para = cell.paragraphs[0]
    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_para.add_run('\n\n\n\n\n[Wklej tutaj screenshot klawiatury]\n\n\n\n\n')
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.size = Pt(12)

    doc.add_paragraph()

    doc.add_page_break()

    # ========== SEKCJA 5: ZAPISYWANIE ==========
    doc.add_heading('5. ZAPISYWANIE I ZATWIERDZANIE', level=1)

    doc.add_paragraph(
        'Po wprowadzeniu wszystkich danych należy je zapisać w systemie.'
    )

    doc.add_paragraph()

    save_steps = doc.add_paragraph()
    save_steps.add_run('Aby zapisać dane:').bold = True

    save_list = [
        'Sprawdź czy wszystkie pola są wypełnione',
        'Kliknij przycisk "Zapisz" lub "Zatwierdź"',
        'Poczekaj na potwierdzenie zapisu',
        'System wyświetli komunikat o powodzeniu',
    ]

    for i, step in enumerate(save_list, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_paragraph()

    # Miejsce na screenshot
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[MIEJSCE NA SCREENSHOT - PRZYCISK ZAPISZ]')
    run.bold = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell_para = cell.paragraphs[0]
    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_para.add_run('\n\n\n\n\n[Wklej tutaj screenshot przycisku zapisz]\n\n\n\n\n')
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.size = Pt(12)

    doc.add_paragraph()

    doc.add_page_break()

    # ========== SEKCJA 6: PROBLEMY ==========
    doc.add_heading('6. ROZWIĄZYWANIE PROBLEMÓW', level=1)

    problems = [
        {
            'problem': 'Waga nie pobiera się automatycznie',
            'rozwiazanie': [
                'Sprawdź połączenie kabla między wagą a komputerem',
                'Zrestartuj aplikację',
                'W razie potrzeby wprowadź wagę ręcznie',
            ]
        },
        {
            'problem': 'Wyświetla się błędna wartość',
            'rozwiazanie': [
                'Sprawdź czy waga jest wyzerowana (tara)',
                'Zdejmij towar i połóż ponownie',
                'Skontaktuj się z przełożonym',
            ]
        },
        {
            'problem': 'Ekran nie reaguje na dotyk',
            'rozwiazanie': [
                'Wyczyść ekran suchą szmatką',
                'Zrestartuj urządzenie',
                'Zgłoś problem do IT',
            ]
        },
        {
            'problem': 'System się zawiesił',
            'rozwiazanie': [
                'Poczekaj 30 sekund',
                'Jeśli nie pomoże - zrestartuj aplikację',
                'W ostateczności zrestartuj komputer/tablet',
            ]
        },
    ]

    for item in problems:
        p = doc.add_paragraph()
        run = p.add_run(f"Problem: {item['problem']}")
        run.bold = True
        run.font.color.rgb = RGBColor(192, 0, 0)

        p2 = doc.add_paragraph()
        p2.add_run('Rozwiązanie:').bold = True

        for sol in item['rozwiazanie']:
            doc.add_paragraph(f'• {sol}')

        doc.add_paragraph()

    doc.add_paragraph()

    # Miejsce na dodatkowe screenshoty
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[MIEJSCE NA DODATKOWE SCREENSHOTY]')
    run.bold = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell_para = cell.paragraphs[0]
    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_para.add_run('\n\n\n\n\n[Wklej tutaj dodatkowe screenshoty]\n\n\n\n\n')
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.size = Pt(12)

    doc.add_page_break()

    # ========== NOTATKI ==========
    doc.add_heading('NOTATKI', level=1)

    doc.add_paragraph('Miejsce na własne notatki i uwagi:')

    doc.add_paragraph()
    doc.add_paragraph('_' * 70)
    doc.add_paragraph()
    doc.add_paragraph('_' * 70)
    doc.add_paragraph()
    doc.add_paragraph('_' * 70)
    doc.add_paragraph()
    doc.add_paragraph('_' * 70)
    doc.add_paragraph()
    doc.add_paragraph('_' * 70)
    doc.add_paragraph()
    doc.add_paragraph('_' * 70)
    doc.add_paragraph()
    doc.add_paragraph('_' * 70)
    doc.add_paragraph()
    doc.add_paragraph('_' * 70)

    doc.add_paragraph()
    doc.add_paragraph()

    # Stopka z kontaktem
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('W razie problemów skontaktuj się z działem IT lub przełożonym')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Zapisz dokument
    output_path = '/home/user/Kalendarz1/Instrukcja_Obslugi_Wagi.docx'
    doc.save(output_path)
    print(f'Instrukcja została zapisana: {output_path}')
    return output_path

if __name__ == '__main__':
    create_instrukcja()
